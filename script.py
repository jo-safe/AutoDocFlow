# -*- coding: utf-8 -*-
import sys, os, re, traceback, json, openpyxl, fitz, copy, tempfile, smtplib
from re import A
from PIL import Image
from docx2pdf import convert
from docx import Document
from mailmerge import MailMerge
from email.message import EmailMessage
from pathlib import Path
from datetime import datetime

def load_excel_data(filePath):
    wb = openpyxl.load_workbook(filePath, data_only=True)
    sheet = wb["Операции"]
    data = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if any(cell is not None for cell in row):
            data.append(list(row))
    return data

def get_contractors(data):
    seen = set()
    unique_values = []
    for row in data:
        if row:
            first_val = row[0]
            if first_val not in seen:
                seen.add(first_val)
                unique_values.append(first_val)
    return unique_values
    
def filter_contractors(data, contractor):
    return [row for row in data if row[0] == contractor]
    
def filter_date_range(data, start, end):
    result = []
    for row in data:
        try:
            date = datetime.strptime(row[5], "%d.%m.%Y")
            if start <= date <= end:
                result.append(row)
        except (IndexError, ValueError):
            # Неверный формат даты/строки
            continue
    return result

def get_dates(data):
    dates = [row[5] for row in data]
    return (min(dates), max(dates))

def get_document_data(data, orgName):
    sum_ = 0
    tableData = {}
    index = 0
    for item in data:
        tableData["operationDate" + str(index)] = item[5].strftime("%d.%m.%Y")
        tableData["operationType" + str(index)] = item[7]
        tableData["operationID" + str(index)] = "(" + str(item[4]) + ")"
        if item[7] == "Поступление средств":
            tableData["incoming" + str(index)] = str(item[6]) + " руб."
        else:
            tableData["outgoing" + str(index)] = str(item[6]) + " руб."
        
        sum_ += (-1 if item[7] == "Поступление средств" else 1) * item[6]
        index += 1
    return (tableData, orgName if sum_ <= 0 else data[0][0], abs(sum_), data[0][3])

def fill_word_template(templatePath, outputPath, mergeFields: dict, tableData: dict, hasTable = True):
    merged_data = {**mergeFields, **tableData}
    
    folder = os.path.dirname(outputPath)
    if folder and not os.path.exists(folder):
        os.makedirs(folder)
        
    with MailMerge(templatePath) as document:
        document.merge(**{k: str(v) for k, v in merged_data.items()})
        document.write(outputPath)
    
    doc = Document(outputPath)
    if hasTable:
        table = doc.tables[0]
        for row_idx in range(len(table.rows) - 1, len(tableData) // 4 + 1, -1):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)
    doc.save(outputPath)
    
def insert_image(pdfPath: str, markerText, imagePath, factor):
    doc = fitz.open(pdfPath)
    for page in doc:
        areas = page.search_for(markerText)
        if not areas:
            continue
        area = areas[0]
        
        page.add_redact_annot(area)
        page.apply_redactions()
        
        with Image.open(imagePath) as img:
            imgWidth, imgHeight = img.size
            
        imgWidth_pt = imgWidth * 72 / img.info.get('dpi', (72, 72))[0] // factor
        imgHeight_pt = imgHeight * 72 / img.info.get('dpi', (72, 72))[1] // factor
        
        x0, y0 = area.x0, area.y0
        rect = fitz.Rect(x0, y0, x0 + imgWidth_pt, y0 + imgHeight_pt)

        page.insert_image(rect, filename=imagePath, overlay=True)
    doc.save(pdfPath[:-4] + "_.pdf")
    return pdfPath[:-4] + "_.pdf"

def get_smtp_settings(email: str):
    domain = email.split('@')[-1].lower()

    settings = {
        'gmail.com': ('smtp.gmail.com', 465),
        'yahoo.com': ('smtp.mail.yahoo.com', 465),
        'yandex.ru': ('smtp.yandex.ru', 465),
        'mail.ru': ('smtp.mail.ru', 465),
        'outlook.com': ('smtp.office365.com', 587),
        'hotmail.com': ('smtp.office365.com', 587)
    }

    if domain in settings:
        return settings[domain]
    else:
        raise ValueError(f"Неизвестный почтовый домен: {domain}")

def send_email(docxTemplatePath, pdfPath, senderEmail, senderPassword, recipientEmail, subject="Документ"):
    doc = Document(docxTemplatePath)
    body_text = "\n".join([para.text for para in doc.paragraphs])
    
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = senderEmail
    msg['To'] = recipientEmail
    msg.set_content(body_text)
    
    pdf_file = Path(pdfPath)
    with pdf_file.open('rb') as f:
        pdf_data = f.read()
        msg.add_attachment(pdf_data,
                           maintype='application',
                           subtype='pdf',
                           filename=pdf_file.name)
    smtp_server, port = get_smtp_settings(senderEmail)
    with smtplib.SMTP_SSL(smtp_server, port) as smtp:
        smtp.login(senderEmail, senderPassword)
        smtp.send_message(msg)
        
def cleanup_similar_pdfs(pdfPath):
    dir_path, target_name = os.path.split(pdfPath)
    prefix = target_name[:-6]

    for fname in os.listdir(dir_path):
        fpath = os.path.join(dir_path, fname)
        if (
            fname.endswith(".pdf") and
            fname != os.path.basename(pdfPath) and
            fname.startswith(prefix)
        ):
            os.remove(fpath)
    newPdfPath = re.sub(r'_+\.pdf$', '.pdf', pdfPath)
    os.rename(pdfPath, os.path.join(os.path.dirname(pdfPath), newPdfPath))
    return newPdfPath
            
def main():
    args = sys.argv[1:]
    
    if (args[0] == "GENERATE"):
        dbPath = args[1]
        docTempPath = args[2]
        signaturePath = args[3]
        stampPath = args[4]
        mailTempPath = args[5]
        outputFolder = args[6]
        orgName = args[7]
        orgPersonName = args[8]
        orgEmail = args[9]
        orgEmailPassword = args[10]
        
        filterPeriod = args[11].lower() == "true"
        startDate = datetime.strptime(args[12], "%d.%m.%Y").date()
        endDate = datetime.strptime(args[13], "%d.%m.%Y").date()

        filterContractor = args[14].lower() == "true"
        contractor = args[15]

        addSignature = args[16].lower() == "true"
        addStamp = args[17].lower() == "true"
        sendEmail = args[18].lower() == "true"
        
        data = load_excel_data(dbPath)
        if filterPeriod: 
            data = filter_date_range(data, startDate, endDate)
        if (not filterContractor) or contractor == "Все контрагенты":
            contractors = get_contractors(data)
        else:
            contractors = [contractor]
        for contractor in contractors:
            iterData = filter_contractors(data, contractor)
            if not filterPeriod:
                (startDate, endDate) = get_dates(iterData)
            iterData.sort(key=lambda x: x[5])
            outputPath = outputFolder + "\\Акт_сверки_" + contractor.replace(' ', '_').replace('"', '_').replace("'", '_')
            (tableData, creditorName, sum_, contractorEmail) = get_document_data(iterData, orgName)
            mergeFields = {
                "year" : endDate.year,
                "dateStart" : startDate.strftime("%d.%m.%Y"),
                "dateEnd" : endDate.strftime("%d.%m.%Y"),
                "orgName" : orgName,
                "orgPersonName" : orgPersonName,
                "contractorPersonName" : iterData[0][2],
                "creditorName" : creditorName,
                "sum" : str(sum_) + " руб."
                }
                
            letterPath = outputPath + "_l.docx"
            docxPath = outputPath + ".docx"
            pdfPath = outputPath + ".pdf"
            
            fill_word_template(docTempPath, docxPath, mergeFields, tableData)
            print(contractor + ": создан документ")
            convert(docxPath)
            if addSignature: 
                pdfPath = insert_image(pdfPath, "{{SIGN}}", signaturePath, 13)
                print(contractor + ": добавлена подпись")
            if addStamp: 
                pdfPath = insert_image(pdfPath, "{{STAMP}}", stampPath, 8)
                print(contractor + ": добавлена печать")
            pdfPath = cleanup_similar_pdfs(pdfPath)

            if sendEmail: 
                fill_word_template(mailTempPath, letterPath, mergeFields, tableData, hasTable=False)
                send_email(letterPath, pdfPath, orgEmail, orgEmailPassword, contractorEmail, "Акт сверки " + contractor)
                print(contractor + ": отправлено письмо")
    elif (args[0] == "GET_CONTRACTORS"):
        dbPath = args[1]
        data = load_excel_data(dbPath)
        contractors = get_contractors(data)
        print("\n".join(contractors))
    elif (args[0] == "GET_CONTRACTOR_DATES"):
        dbPath = args[1]
        contractor = args[2]
        data = load_excel_data(dbPath)
        data = filter_contractors(data, contractor)
        (minDate, maxDate) = get_dates(data)
        print("{0} {1}".format(minDate.strftime("%d.%m.%Y"),  maxDate.strftime("%d.%m.%Y")).replace("\r", ""))

if __name__ == "__main__":
    main()